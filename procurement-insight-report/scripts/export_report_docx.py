#!/usr/bin/env python3
"""Export a Procurement MCP report JSON payload into a DOCX report."""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


CONFIDENCE_COLORS = {
    "high": RGBColor(22, 101, 52),
    "mixed": RGBColor(180, 83, 9),
    "low": RGBColor(185, 28, 28),
}


def load_report(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_doc(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for style_name, size in [("Title", 24), ("Heading 1", 16), ("Heading 2", 13)]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)


def add_title_block(document: Document, report: dict) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(report.get("title", "Procurement Report"))
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(24)

    if report.get("subtitle"):
        subtitle = document.add_paragraph(report["subtitle"])
        subtitle.runs[0].font.name = "Arial"
        subtitle.runs[0].font.size = Pt(11)

    meta = document.add_paragraph()
    meta.add_run("Reporting window: ").bold = True
    meta.add_run(report.get("reporting_window", ""))

    confidence = str(report.get("confidence", "mixed")).lower()
    conf_para = document.add_paragraph()
    conf_run = conf_para.add_run(f"Confidence: {confidence.capitalize()}")
    conf_run.bold = True
    conf_run.font.color.rgb = CONFIDENCE_COLORS.get(confidence, CONFIDENCE_COLORS["mixed"])


def add_bullets(document: Document, heading: str, items: list[str]) -> None:
    document.add_heading(heading, level=1)
    if not items:
        document.add_paragraph("No items provided.")
        return
    for item in items:
        document.add_paragraph(str(item), style="List Bullet")


def add_kpi_table(document: Document, kpis: list[dict]) -> None:
    if not kpis:
        return
    document.add_heading("KPI Summary", level=1)
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    set_cell_shading(hdr[0], "DFF5F2")
    set_cell_shading(hdr[1], "DFF5F2")
    for row in kpis:
        cells = table.add_row().cells
        cells[0].text = str(row.get("label", "Metric"))
        cells[1].text = str(row.get("value", ""))


def add_chart_tables(document: Document, charts: list[dict]) -> None:
    if not charts:
        return
    document.add_heading("Chart Data", level=1)
    for chart in charts:
        document.add_heading(chart.get("title", "Chart"), level=2)
        series = chart.get("series", [])
        if not series:
            document.add_paragraph("No chart data provided.")
            continue
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Label"
        hdr[1].text = "Value"
        set_cell_shading(hdr[0], "F3EDE2")
        set_cell_shading(hdr[1], "F3EDE2")
        for point in series:
            row = table.add_row().cells
            row[0].text = str(point.get("label", ""))
            row[1].text = str(point.get("value", ""))


def build_docx(report: dict, output_path: Path) -> None:
    document = Document()
    style_doc(document)
    add_title_block(document, report)
    add_bullets(document, "Executive Summary", report.get("summary", []))
    add_kpi_table(document, report.get("kpis", []))
    add_chart_tables(document, report.get("charts", []))
    add_bullets(document, "Facts", report.get("facts", []))
    add_bullets(document, "Insights", report.get("insights", []))
    add_bullets(document, "Recommended Actions", report.get("recommended_actions", []))
    add_bullets(document, "Confidence and Caveats", report.get("caveats", []))
    if report.get("sources"):
        add_bullets(document, "Sources", report.get("sources", []))
    document.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_json", type=Path, help="Path to the normalized report JSON file.")
    parser.add_argument("output_docx", type=Path, help="Path to write the DOCX report.")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    report = load_report(args.input_json)
    build_docx(report, args.output_docx)
    print(f"Rendered DOCX report to {args.output_docx}")


if __name__ == "__main__":
    main()
